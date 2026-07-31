"""巡检 DOCX 的离线结构提取与安全检查。"""
from __future__ import annotations

import re
import zipfile
import shutil
from copy import deepcopy
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from docx import Document
from fastapi import HTTPException

IP_RE = re.compile(r"(?<!\d)(?:25[0-5]|2[0-4]\d|1?\d?\d)(?:\.(?:25[0-5]|2[0-4]\d|1?\d?\d)){3}(?!\d)")
READ_ONLY_PREFIXES = ("display ", "dis ", "show ")
BLOCKED_WORDS = ("reboot", "restart", "reset", "save", "delete", "undo ", "shutdown", "format", "configure", "system-view", "commit")
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
MAX_DOCX_FILES = 2000
MAX_DOCX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024

def inspect_docx_security(path: Path) -> None:
    if path.suffix.lower() != ".docx" or not zipfile.is_zipfile(path):
        raise HTTPException(400, "仅支持有效的 .docx 文件")
    with zipfile.ZipFile(path) as archive:
        members = archive.infolist(); names = [member.filename for member in members]
        if len(members) > MAX_DOCX_FILES or sum(member.file_size for member in members) > MAX_DOCX_UNCOMPRESSED_BYTES:
            raise HTTPException(400, "DOCX 解压后的大小或文件数超出限制")
        if any(Path(name.replace("\\", "/")).is_absolute() or ".." in Path(name.replace("\\", "/")).parts for name in names):
            raise HTTPException(400, "DOCX 包含非法内部路径")
        if any(name.startswith("word/embeddings/") or name.endswith("vbaProject.bin") for name in names):
            raise HTTPException(400, "DOCX 包含嵌入对象或宏，已拒绝")
        for name in names:
            if not name.endswith(".rels"): continue
            root = ET.fromstring(archive.read(name))
            for rel in root.findall(f"{{{REL_NS}}}Relationship"):
                if rel.attrib.get("TargetMode") == "External":
                    raise HTTPException(400, "DOCX 包含外部关系，已拒绝")

def validate_read_only_command(command: str) -> str:
    normalized = " ".join(command.strip().split())
    lower = normalized.lower()
    if not lower.startswith(READ_ONLY_PREFIXES) or any(word in lower for word in BLOCKED_WORDS):
        raise HTTPException(400, f"命令不是允许的只读命令: {normalized}")
    if any(ch in normalized for ch in ("\r", "\n", ";", "&", "|")):
        raise HTTPException(400, f"命令包含非法控制符: {normalized}")
    return normalized

def parse_template(path: Path, *, system_key: str, display_name: str) -> dict[str, Any]:
    inspect_docx_security(path)
    doc = Document(path)
    devices: list[dict[str, Any]] = []
    non_commands: list[dict[str, Any]] = []
    for table_index, table in enumerate(doc.tables):
        header_text = "\n".join(cell.text for row in table.rows[:3] for cell in row.cells)
        command_header_row, command_column, result_column = _find_command_columns(table)
        # 设备清单、版本信息等辅助表也可能包含名称/IP，但没有巡检命令列，不能当作设备表。
        if command_column is None:
            continue
        labeled_name = _extract_labeled_value(table, ("设备名称", "网元名称", "主机名", "名称"))
        prompt_name = _extract_prompt_name(table, command_header_row, result_column)
        # 历史报告的表头可能写错，命令结果开头的真实设备提示符优先级更高。
        name = prompt_name or labeled_name or _extract_name(header_text)
        ip = _extract_labeled_value(table, ("设备IP", "网元IP", "管理IP", "IP"))
        ips = IP_RE.findall(header_text)
        if not ip and ips:
            ip = ips[0]
        commands: list[dict[str, Any]] = []
        for row_index, row in enumerate(table.rows):
            if row_index <= command_header_row or command_column >= len(row.cells):
                continue
            command_text = row.cells[command_column].text.strip()
            if command_text.lower().startswith(READ_ONLY_PREFIXES):
                command = validate_read_only_command(command_text)
                commands.append({"command": command, "timeout_seconds": 120, "result_cell": {"table": table_index, "row": row_index, "column": result_column if result_column is not None else min(command_column + 1, len(row.cells) - 1)}})
            elif command_text:
                label_column = 1 if len(row.cells) > 1 and command_column != 1 else 0
                label = row.cells[label_column].text.strip()
                if label:
                    non_commands.append({"label": label, "result_cell": {"table": table_index, "row": row_index, "column": result_column if result_column is not None else min(command_column + 1, len(row.cells) - 1)}, "mode": "preserve", "value": "", "source_command": ""})
        devices.append({"order": len(devices) + 1, "name": name or f"设备{len(devices)+1}", "ip": ip, "driver": _guess_driver(commands), "commands": commands, "table_index": table_index})
    if not devices: raise HTTPException(400, "未在 DOCX 表格中识别到设备")
    return {"system_key": system_key, "display_name": display_name, "template": path.name, "devices": devices, "non_command_rules": non_commands}

def _find_command_columns(table) -> tuple[int, int | None, int | None]:
    """返回命令表头行及命令、结果列，兼容不同模板的列顺序。"""
    for row_index, row in enumerate(table.rows[:8]):
        headers = [" ".join(cell.text.split()).casefold() for cell in row.cells]
        command_column = next((i for i, value in enumerate(headers) if "巡检命令" in value or value in {"命令", "检查命令"}), None)
        if command_column is None:
            continue
        result_column = next((i for i, value in enumerate(headers) if "巡检结果" in value or value in {"结果", "检查结果"}), None)
        return row_index, command_column, result_column
    # 旧模板没有表头：仅在单列确实包含只读命令时启用，避免把普通清单误判为设备表。
    command_counts: dict[int, int] = {}
    max_columns = 0
    for row in table.rows:
        max_columns = max(max_columns, len(row.cells))
        for column, cell in enumerate(row.cells):
            if cell.text.strip().lower().startswith(READ_ONLY_PREFIXES):
                command_counts[column] = command_counts.get(column, 0) + 1
    if command_counts:
        command_column = max(command_counts, key=command_counts.get)
        result_column = command_column + 1 if command_column + 1 < max_columns else None
        return -1, command_column, result_column
    return -1, None, None

def _extract_labeled_value(table, labels: tuple[str, ...]) -> str:
    wanted = {label.casefold() for label in labels}
    for row in table.rows[:4]:
        for index, cell in enumerate(row.cells):
            if " ".join(cell.text.split()).casefold() in wanted and index + 1 < len(row.cells):
                value = row.cells[index + 1].text.strip()
                if value:
                    return value
    return ""

def _extract_prompt_name(table, header_row: int, result_column: int | None) -> str:
    if result_column is None:
        return ""
    for row in table.rows[header_row + 1:]:
        if result_column >= len(row.cells):
            continue
        match = re.match(r"\s*<([^<>\r\n]+)>\s*(?:display|dis|show)\b", row.cells[result_column].text, re.I)
        if match:
            return match.group(1).strip()
    return ""

def _extract_name(text: str) -> str:
    for pattern in (r"(?:设备名称|名称|主机名|型号)\s*[:：]\s*([^\n\r]+)", r"<([^<>\s]+)>"):
        match = re.search(pattern, text, re.I)
        if match: return match.group(1).strip()
    return ""

def _guess_driver(commands: list[dict[str, Any]]) -> str:
    return "generic_show" if commands and all(item["command"].lower().startswith("show ") for item in commands) else "huawei_vrp"

def validate_snapshot(snapshot: dict[str, Any]) -> dict[str, Any]:
    devices = snapshot.get("devices")
    if not isinstance(devices, list) or not devices or len(devices) > 100: raise HTTPException(400, "设备数量必须为 1 到 100")
    try: devices.sort(key=lambda item: int(item.get("order", 0)))
    except (TypeError, ValueError): raise HTTPException(400, "设备顺序必须是整数") from None
    names: set[str] = set(); ips: set[str] = set()
    for order, device in enumerate(devices, 1):
        name = str(device.get("name", "")).strip(); ip = str(device.get("ip", "")).strip()
        if not name: raise HTTPException(400, f"第 {order} 台设备名称为空")
        name_key = name.casefold()
        if name_key in names: raise HTTPException(409, f"设备名称冲突: {name}")
        names.add(name_key)
        if ip:
            if not IP_RE.fullmatch(ip): raise HTTPException(400, f"管理 IP 不合法: {ip}")
            if ip in ips: raise HTTPException(409, f"管理 IP 冲突: {ip}")
            ips.add(ip)
        driver = device.get("driver", "huawei_vrp")
        if driver not in {"huawei_vrp", "generic_show"}: raise HTTPException(400, f"未知驱动: {driver}")
        raw_commands = device.get("commands", [])
        if not isinstance(raw_commands, list) or not raw_commands: raise HTTPException(400, f"设备 {name} 没有命令")
        for item in raw_commands:
            item["command"] = validate_read_only_command(str(item.get("command", "")))
            timeout = int(item.get("timeout_seconds", 120))
            if timeout < 1 or timeout > 900: raise HTTPException(400, f"命令超时必须在 1 到 900 秒")
            item["timeout_seconds"] = timeout
            mapping = item.get("result_cell")
            if not isinstance(mapping, dict) or not all(isinstance(mapping.get(key), int) and mapping[key] >= 0 for key in ("table", "row", "column")):
                raise HTTPException(400, f"命令 {item['command']} 的回填位置不合法")
        device["order"] = order
    rules = snapshot.get("non_command_rules", [])
    if not isinstance(rules, list): raise HTTPException(400, "非命令规则必须是数组")
    for rule in rules:
        if rule.get("mode") not in {"preserve", "fixed", "derived", "manual"}: raise HTTPException(400, "非命令规则模式不合法")
        mapping = rule.get("result_cell")
        if not isinstance(mapping, dict) or not all(isinstance(mapping.get(key), int) and mapping[key] >= 0 for key in ("table", "row", "column")):
            raise HTTPException(400, "非命令规则回填位置不合法")
        if rule["mode"] == "derived" and not str(rule.get("source_command", "")).strip(): raise HTTPException(400, "命令推导规则缺少来源命令")
    return snapshot


def validate_snapshot_against_template(snapshot: dict[str, Any], template_path: Path) -> dict[str, Any]:
    """校验设备、命令和回填坐标确实对应当前 DOCX。"""
    if template_path.suffix.lower() != ".docx" or not template_path.is_file():
        raise HTTPException(409, f"模板文件不存在或不是 DOCX: {template_path.name}")
    doc = Document(template_path)
    occupied: set[tuple[int, int, int]] = set()

    def command_key(value: str) -> str:
        normalized = "".join(character for character in value.casefold() if character.isalnum())
        for prefix in ("display", "dis", "show"):
            if normalized.startswith(prefix):
                return normalized[len(prefix):]
        return normalized

    for device in snapshot.get("devices", []):
        name = str(device.get("name", "")).strip() or "未命名设备"
        table_index = device.get("table_index")
        if not isinstance(table_index, int) or not 0 <= table_index < len(doc.tables):
            raise HTTPException(409, f"设备 {name} 的表格索引无效: {table_index}")
        table = doc.tables[table_index]
        for item in device.get("commands", []):
            command = str(item.get("command", "")).strip()
            mapping = item.get("result_cell", {})
            row_index, column_index = mapping.get("row"), mapping.get("column")
            if mapping.get("table") != table_index:
                raise HTTPException(409, f"设备 {name} 的命令 {command} 绑定到了其他表格")
            if (
                not isinstance(row_index, int)
                or not isinstance(column_index, int)
                or not 0 <= row_index < len(table.rows)
                or not 0 <= column_index < len(table.rows[row_index].cells)
            ):
                raise HTTPException(409, f"设备 {name} 的命令 {command} 回填坐标超出模板范围")
            location = (table_index, row_index, column_index)
            if location in occupied:
                raise HTTPException(409, f"设备 {name} 的命令 {command} 与其他项目使用了同一回填单元格")
            occupied.add(location)
            expected = command_key(command)
            row_commands = {
                command_key(cell.text.strip())
                for index, cell in enumerate(table.rows[row_index].cells)
                if index != column_index and cell.text.strip().casefold().startswith(READ_ONLY_PREFIXES)
            }
            if expected not in row_commands:
                raise HTTPException(409, f"设备 {name} 的命令 {command} 与模板第 {row_index + 1} 行不一致")

    for rule in snapshot.get("non_command_rules", []):
        mapping = rule.get("result_cell", {})
        table_index, row_index, column_index = (
            mapping.get("table"), mapping.get("row"), mapping.get("column")
        )
        if (
            not isinstance(table_index, int)
            or not isinstance(row_index, int)
            or not isinstance(column_index, int)
            or not 0 <= table_index < len(doc.tables)
            or not 0 <= row_index < len(doc.tables[table_index].rows)
            or not 0 <= column_index < len(doc.tables[table_index].rows[row_index].cells)
        ):
            raise HTTPException(409, f"非命令项 {rule.get('label', '')} 回填坐标超出模板范围")
        location = (table_index, row_index, column_index)
        if location in occupied:
            raise HTTPException(409, f"非命令项 {rule.get('label', '')} 与其他项目使用了同一回填单元格")
        occupied.add(location)
    return snapshot

def merge_incremental_template(base_path: Path, added_path: Path, output_path: Path) -> int:
    """把增量文档中的设备表格追加到旧模板，保留原表格 OOXML 与样式。"""
    base = Document(base_path); added = Document(added_path)
    if len(base.sections) != len(added.sections): raise HTTPException(409, "增量模板版式不兼容，请使用完整替换")
    for left, right in zip(base.sections, added.sections):
        geometry = (left.page_width, left.page_height, left.left_margin, left.right_margin, left.orientation)
        other = (right.page_width, right.page_height, right.left_margin, right.right_margin, right.orientation)
        if geometry != other: raise HTTPException(409, "增量模板版式不兼容，请使用完整替换")
    offset = len(base.tables)
    for table in added.tables:
        base.element.body.insert(-1, deepcopy(table._element))
    temporary = output_path.with_suffix(".merged.docx"); base.save(temporary); shutil.move(temporary, output_path)
    return offset

def shift_table_mappings(snapshot: dict[str, Any], offset: int) -> None:
    for device in snapshot.get("devices", []):
        if "table_index" in device: device["table_index"] += offset
        for command in device.get("commands", []): command["result_cell"]["table"] += offset
    for rule in snapshot.get("non_command_rules", []): rule["result_cell"]["table"] += offset
