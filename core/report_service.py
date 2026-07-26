from __future__ import annotations

import json
import pickle
import re
from collections import Counter
from concurrent.futures import ProcessPoolExecutor, as_completed
from concurrent.futures.process import BrokenProcessPool
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable

from docx import Document

from core import docx_engine as engine

RE_IP = re.compile(r"(\d+\.\d+\.\d+\.[1-9]\d*)")
RE_DATE = re.compile(r"\d{4}[-]?\d{2}[-]?\d{2}|\d{2}月\d{2}日")
# 只认行首的提示符形态，避免抓到正文里的 <...> 垃圾串
RE_PROMPT_HOST = re.compile(r"^(?:HRP_[MS])?<([^<>\s]+)>", re.M)
ROOM_ENV_LABEL = "机房电源、空调、温湿度检查"
ROOM_ENV_ALARM_KEYWORDS = (
    "power",
    "pwr",
    "fan",
    "temperature",
    "humid",
    "humidity",
    "environment",
    "电源",
    "风扇",
    "温度",
    "湿度",
    "环境",
)
ROOM_ENV_EN_PATTERNS = tuple(
    re.compile(pattern, re.I)
    for pattern in (
        r"\bpower\b",
        r"\bpwr\d*\b",
        r"\bfan\d*\b",
        r"\btemperature\b",
        r"\bhumid(?:ity)?\b",
        r"\benvironment\b",
    )
)
ROOM_ENV_BAD_KEYWORDS = (
    "critical",
    "major",
    "minor",
    "fault",
    "abnormal",
    "absent",
    "unregistered",
    "shutdown",
    "failed",
    "failure",
    "off",
    "not present",
    "lost",
    "异常",
    "故障",
    "关闭",
    "离线",
)
GPRS_CONTACT_LINES = (
    "维护人员 :苗向-18338289920",
    "负责人员-对接人员:费泽庆",
    "李诚-13648914080",
)
GPRS_CONTACT_MARKERS = {
    "维护人员",
    "维护人员 :苗向-18338289920",
    "杜康-19713769776",
    "负责人员-对接人员",
    "负责人员-对接人员:费泽庆",
    "费泽庆",
    "李诚-13648914080",
}


@dataclass
class ReportPaths:
    root: Path
    config_path: Path
    logs_base: Path
    templates_dir: Path
    output_base: Path


@dataclass
class GenerationSummary:
    target_date: str
    log_root: str
    output_dir: str
    generated_files: list[str]
    audit_lines: list[str]


class LogObject:
    def __init__(self, path: Path, all_config: dict):
        self.filename = path.name
        self.read_error = ""
        self.text = self._read_safe(path)
        self.sections = engine.parse_sections(self.text)
        self.norm_cache = {engine.normalize(cmd): cmd for cmd in self.sections}

        prompts = [name.strip() for name in RE_PROMPT_HOST.findall(self.text) if name.strip()]
        self.real_hostname = Counter(prompts).most_common(1)[0][0] if prompts else ""

        ips_in_name = RE_IP.findall(self.filename)
        self.filename_ip = next(
            (ip for ip in ips_in_name if not ip.startswith(("127.", "0.", "1.23.", "1.3."))),
            "",
        )
        self.content_ips = set(RE_IP.findall(self.text))

        stem = path.stem
        parts = stem.split("_")
        prefix = parts[0].upper() if parts else ""
        is_sys_prefix = any(key.upper() in prefix for key in list(all_config.keys()) + ["NETMGMT", "NM", "SMS", "GPRS"])
        self.file_subject = (parts[1] if len(parts) >= 2 and is_sys_prefix else parts[0]).lower()

    def _read_safe(self, path: Path) -> str:
        for encoding in ("utf-8", "gbk", "gb18030"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
            except OSError as error:
                self.read_error = str(error)
                return ""
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except OSError as error:
            self.read_error = str(error)
            return ""


def default_paths(root: Path | None = None) -> ReportPaths:
    if root is None:
        root = Path(__file__).resolve().parents[1]
    return ReportPaths(
        root=root,
        config_path=root / "config" / "report.json",
        logs_base=root / "Logs_待处理",
        templates_dir=root / "assets" / "templates",
        output_base=root / "巡检报告",
    )


def load_config(config_path: Path) -> dict:
    with config_path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def find_latest_log_dir(logs_base: Path) -> Path:
    date_dirs = sorted(path for path in logs_base.iterdir() if path.is_dir() and "-" in path.name)
    if not date_dirs:
        raise FileNotFoundError(f"未在 {logs_base} 发现日期目录")
    return date_dirs[-1]


def find_section_output(sections: dict[str, str], cache: dict[str, str], *keywords: str) -> str:
    for command, output in sections.items():
        normalized = cache.get(engine.normalize(command), command)
        target = engine.normalize(normalized)
        if all(engine.normalize(keyword) in target for keyword in keywords):
            return output
    return ""


def resolve_room_environment_status(log: LogObject) -> str:
    device_output = find_section_output(log.sections, log.norm_cache, "device")
    logbuffer_output = find_section_output(log.sections, log.norm_cache, "logbuffer")
    device_text = device_output.lower()
    logbuffer_text = logbuffer_output.lower()

    if not device_output and not logbuffer_output:
        return "未采集到相关日志，请人工核查。"

    device_has_good_signal = False
    device_bad_lines: list[str] = []
    if device_output:
        if "alarm" in device_text:
            bad_words = ("critical", "major", "minor", "abnormal", "fault")
            device_has_good_signal = "normal" in device_text and not any(bad in device_text for bad in bad_words)
        else:
            bad_words = ("fault", "abnormal", "absent", "unregistered")
            lines = [line.strip().lower() for line in device_output.splitlines() if line.strip()]
            device_has_good_signal = bool(lines) and all(
                not any(bad in line for bad in bad_words)
                for line in lines
                if line[:1].isdigit()
            )
        if not device_has_good_signal:
            for line in device_output.splitlines():
                if not any(bad in line.lower() for bad in bad_words):
                    continue
                device_bad_lines.append(re.sub(r"\s+", " ", line.strip()))
                if len(device_bad_lines) >= 2:
                    break

    abnormal_hits: list[str] = []
    for line in logbuffer_output.splitlines():
        normalized_line = line.lower()
        has_env_keyword = any(keyword in normalized_line for keyword in ROOM_ENV_ALARM_KEYWORDS) or any(
            pattern.search(line) for pattern in ROOM_ENV_EN_PATTERNS
        )
        if not has_env_keyword:
            continue
        if any(bad in normalized_line for bad in ROOM_ENV_BAD_KEYWORDS):
            abnormal_hits.append(line.strip())
            if len(abnormal_hits) >= 2:
                break

    if abnormal_hits:
        return "；".join(abnormal_hits)
    if device_output and not device_has_good_signal:
        if device_bad_lines:
            return "设备状态异常：" + "；".join(device_bad_lines) + "，请人工核查。"
        return "设备状态异常，请人工核查。"
    return "正常。"


def safe_update_paragraph(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.text = str(new_text)
        return
    paragraph.runs[0].text = str(new_text)
    for run in paragraph.runs[1:]:
        run.text = ""


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def ensure_gprs_contacts(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    anchor_index = next((index for index, p in enumerate(paragraphs) if p.text.strip() == "维护人员"), None)

    if anchor_index is None:
        return

    safe_update_paragraph(paragraphs[anchor_index], GPRS_CONTACT_LINES[0])

    current_index = anchor_index + 1
    for line in GPRS_CONTACT_LINES[1:]:
        while current_index < len(paragraphs) and not paragraphs[current_index].text.strip():
            current_index += 1
        if current_index < len(paragraphs):
            safe_update_paragraph(paragraphs[current_index], line)
            current_index += 1
        else:
            doc.add_paragraph(line)

    removable: list = []
    while current_index < len(paragraphs):
        text = paragraphs[current_index].text.strip()
        if not text:
            break
        if text not in GPRS_CONTACT_MARKERS:
            break
        removable.append(paragraphs[current_index])
        current_index += 1

    for paragraph in removable:
        remove_paragraph(paragraph)


def fill_gprs_contact_rows(table) -> None:
    for row in table.rows:
        cells = row.cells
        if not cells:
            continue
        labels = [cell.text.strip() for cell in cells]
        for index, label in enumerate(labels):
            if "维护人员" in label:
                target_index = min(index + 1, len(cells) - 1)
                engine.set_cell_text(cells[target_index], "苗向-18338289920")
                break
            if "负责人员-对接人员" in label or ("负责人员" in label and "对接人员" in label):
                target_index = min(index + 1, len(cells) - 1)
                engine.set_cell_text(cells[target_index], "费泽庆\n李诚-13648914080")
                break


def is_gprs_contact_row(cells) -> bool:
    for cell in cells:
        text = cell.text.strip()
        if "维护人员" in text:
            return True
        if "负责人员-对接人员" in text or ("负责人员" in text and "对接人员" in text):
            return True
    return False


def find_match(log_pool: list[LogObject], target_ip: str, template_host: str, config_host: str) -> tuple[LogObject | None, str]:
    template_host = template_host.lower()
    config_host = config_host.lower()

    if target_ip:
        for index, log in enumerate(log_pool):
            if target_ip == log.filename_ip:
                return log_pool.pop(index), f"文件同IP({target_ip})"

    for target in (config_host, template_host):
        if not target or len(target) < 3:
            continue
        for index, log in enumerate(log_pool):
            if target == log.real_hostname.lower():
                return log_pool.pop(index), f"设备名全等({log.real_hostname})"

    for target in (config_host, template_host):
        if not target or len(target) < 3:
            continue
        for index, log in enumerate(log_pool):
            subject = log.file_subject
            if target == subject or (target in subject and target[-3:] == subject[-3:]) or (subject in target and subject[-3:] == target[-3:]):
                return log_pool.pop(index), f"业务代号({log.file_subject})"

    if target_ip:
        for index, log in enumerate(log_pool):
            if target_ip in log.content_ips:
                return log_pool.pop(index), f"日志内容IP({target_ip})"

    return None, ""


def process_system(
    sys_key: str,
    sys_info: dict,
    log_root: Path,
    output_dir: Path,
    target_date: str,
    all_configs: dict,
    templates_dir: Path,
) -> tuple[list[str], list[str]]:
    audit_lines = [f"\n=== {sys_info['display_name']} ==="]
    generated_files: list[str] = []
    log_dir = log_root / sys_key
    if not log_dir.exists():
        log_dir = log_root / sys_key.replace("NM", "NetMgmt")
        if not log_dir.exists():
            return audit_lines, generated_files

    log_pool: list[LogObject] = []
    for path in log_dir.glob("*.log"):
        if "summary" in path.name:
            continue
        try:
            log_object = LogObject(path, all_configs)
        except Exception as error:
            audit_lines.append(f"× 日志解析失败 {path.name}: {error}")
            continue
        if log_object.read_error:
            audit_lines.append(f"× 日志读取失败 {path.name}: {log_object.read_error}")
        log_pool.append(log_object)

    template_file = templates_dir / sys_info["template"]
    if not template_file.exists():
        return [f"× {sys_key}: 找不到模板 {template_file.name}"], generated_files

    doc = Document(template_file)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if sys_info.get("is_dual_title") and "网管网络设备" in text:
            safe_update_paragraph(paragraph, sys_info["display_name"])
        elif "日巡检报告" in text:
            safe_update_paragraph(paragraph, f"{sys_info['display_name']}{target_date}日巡检报告")
        elif RE_DATE.search(text):
            safe_update_paragraph(paragraph, RE_DATE.sub(target_date, text))

    for index, table in enumerate(doc.tables, 1):
        if sys_key == "GPRS":
            fill_gprs_contact_rows(table)

        target_ip = ""
        template_host = ""
        for row in table.rows[:2]:
            for cell in row.cells:
                match = RE_IP.search(cell.text)
                if match:
                    target_ip = match.group(1)
                if any(keyword in cell.text for keyword in ["型号", "名称", "设备"]):
                    parts = re.split(r"[:：]", cell.text)
                    template_host = parts[1].strip() if len(parts) > 1 else ""

        config_host = sys_info.get("hosts", {}).get(str(index), "")
        log, reason = find_match(log_pool, target_ip, template_host, config_host)
        if not log:
            audit_lines.append(f"× [{index:02d}] 匹配失败 (IP:{target_ip}, Host:{template_host or config_host})")
            continue

        audit_lines.append(f"√ [{index:02d}] 命中 {log.filename} via {reason}")
        for row_index, row in enumerate(table.rows):
            cells = row.cells
            if sys_key == "GPRS" and is_gprs_contact_row(cells):
                fill_gprs_contact_rows(table)
                continue
            if row_index == 1:
                for cell in cells:
                    if "IP" in cell.text.upper():
                        match = RE_IP.search(cell.text)
                        resolved_ip = match.group(1) if match and not match.group(1).startswith("127.") else log.filename_ip
                        engine.set_cell_text(cell, f"IP: {resolved_ip or '待补充'}")
                    elif RE_DATE.search(cell.text) or cell.text.strip().isdigit():
                        engine.set_cell_text(cell, target_date.replace("-", ""))
            elif row_index == 2:
                for cell in cells[1:]:
                    if not cell.text.strip() or cell.text in ["杜康", "刘关雷", "苗向"]:
                        engine.set_cell_text(cell, "苗向")
            elif row_index >= 4 and len(cells) >= 3:
                wanted = cells[1].text.strip()
                if wanted:
                    if ROOM_ENV_LABEL in wanted:
                        engine.set_cell_text(cells[2], resolve_room_environment_status(log))
                    else:
                        engine.set_cell_text(
                            cells[2],
                            engine.select_command_output(
                                log.sections,
                                wanted,
                                log.norm_cache,
                                system_key=sys_key,
                                template_host=template_host or config_host,
                                target_ip=target_ip,
                            ),
                        )

        if sys_key == "GPRS":
            fill_gprs_contact_rows(table)

    if sys_key == "GPRS":
        ensure_gprs_contacts(doc)

    file_name = f"{(sys_key if sys_info.get('is_english_name') else sys_info['display_name'])}{target_date}日巡检报告.docx"
    output_path = output_dir / file_name
    try:
        doc.save(output_path)
    except OSError as error:
        audit_lines.append(f"× {sys_key}: 报告保存失败 {file_name} ({error})")
        return audit_lines, generated_files
    generated_files.append(str(output_path))
    return audit_lines, generated_files


def generate_reports(
    *,
    paths: ReportPaths,
    log_root: Path | None = None,
    target_date: str | None = None,
    output_dir: Path | None = None,
    max_workers: int | None = None,
    progress_callback: Callable[[int, int, str, dict], None] | None = None,
) -> GenerationSummary:
    configs = load_config(paths.config_path)
    log_root = log_root or find_latest_log_dir(paths.logs_base)
    target_date = target_date or datetime.now().strftime("%Y-%m-%d")
    output_dir = output_dir or (paths.output_base / target_date)
    output_dir.mkdir(parents=True, exist_ok=True)

    all_audit: list[str] = []
    generated_files: list[str] = []
    total_systems = len(configs)
    worker_count = max_workers or min(4, max(1, len(configs)))
    if worker_count <= 1:
        for completed_count, (key, value) in enumerate(configs.items(), 1):
            audit_lines, files = process_system(key, value, log_root, output_dir, target_date, configs, paths.templates_dir)
            all_audit.extend(audit_lines)
            generated_files.extend(files)
            if progress_callback is not None:
                progress_callback(completed_count, total_systems, key, value)
    else:
        # 单个系统的失败由 per-future 兜底，这里只处理进程池本身不可用的情况
        handled_keys: set[str] = set()
        try:
            with ProcessPoolExecutor(max_workers=worker_count) as executor:
                futures = {
                    executor.submit(process_system, key, value, log_root, output_dir, target_date, configs, paths.templates_dir): key
                    for key, value in configs.items()
                }
                for future in as_completed(futures):
                    key = futures[future]
                    try:
                        audit_lines, files = future.result()
                    except BrokenProcessPool:
                        continue
                    except Exception as error:
                        all_audit.append(f"× {key}: 生成失败 ({error})")
                    else:
                        all_audit.extend(audit_lines)
                        generated_files.extend(files)
                    handled_keys.add(key)
                    if progress_callback is not None:
                        progress_callback(len(handled_keys), total_systems, key, configs[key])
        except (OSError, BrokenProcessPool, pickle.PickleError) as error:
            all_audit.append(f"× 进程池不可用，转为串行生成 ({error})")

        for key, value in configs.items():
            if key in handled_keys:
                continue
            try:
                audit_lines, files = process_system(key, value, log_root, output_dir, target_date, configs, paths.templates_dir)
            except Exception as error:
                all_audit.append(f"× {key}: 生成失败 ({error})")
            else:
                all_audit.extend(audit_lines)
                generated_files.extend(files)
            handled_keys.add(key)
            if progress_callback is not None:
                progress_callback(len(handled_keys), total_systems, key, value)

    try:
        (output_dir / "audit_matching_result.txt").write_text("\n".join(all_audit), encoding="utf-8")
    except OSError:
        pass
    return GenerationSummary(
        target_date=target_date,
        log_root=str(log_root),
        output_dir=str(output_dir),
        generated_files=generated_files,
        audit_lines=all_audit,
    )
