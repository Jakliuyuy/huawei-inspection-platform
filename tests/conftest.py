"""共享 fixture 与工具。

黄金基线的核心约束：**不能对 .docx 做二进制哈希**。docx 是 zip，内部
docProps/core.xml 带 <dcterms:modified>，python-docx 每次保存都不同，
二进制哈希 100% 不稳定。这里改为提取"内容指纹"：按序取出全部
paragraph 与 table cell 的文本再哈希。
"""

from __future__ import annotations

import hashlib
import sys
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

# 真实巡检日志语料：139 个设备日志，9 个系统。
# 这是仓库外的本地数据，缺失时相关用例整体跳过而不是失败。
REAL_LOG_ROOT = REPO_ROOT.parent / ".temp_extracted" / "2026-7-24" / "2026-7-24"

# 基线固定用这个日期，避免"今天"漂移导致指纹每天都变。
GOLDEN_DATE = "2026-07-24"

GOLDEN_DIR = Path(__file__).resolve().parent / "golden"


def docx_fingerprint(path: Path) -> str:
    """按序提取 docx 的全部可见文本并哈希。"""
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return hashlib.sha256("\x1f".join(parts).encode("utf-8")).hexdigest()


def run_generation(output_dir: Path, *, only_systems=None, target_date: str = GOLDEN_DATE):
    """跑一次报告生成。

    max_workers=1 强制串行：进程池下 audit 行序按完成顺序 extend，
    多进程时行序不确定，无法用于逐行比对。
    """
    from core.report_service import ReportPaths, generate_reports

    paths = ReportPaths(
        root=REPO_ROOT,
        config_path=REPO_ROOT / "config" / "report.json",
        logs_base=REAL_LOG_ROOT.parent,
        templates_dir=REPO_ROOT / "assets" / "templates",
        output_base=output_dir,
    )
    kwargs = dict(
        paths=paths,
        log_root=REAL_LOG_ROOT,
        target_date=target_date,
        output_dir=output_dir,
        max_workers=1,
    )
    if only_systems is not None:
        kwargs["only_systems"] = only_systems
    return generate_reports(**kwargs)


def fingerprints_of(output_dir: Path) -> dict[str, str]:
    return {p.name: docx_fingerprint(p) for p in sorted(output_dir.glob("*.docx"))}


requires_real_logs = pytest.mark.skipif(
    not REAL_LOG_ROOT.is_dir(),
    reason=f"真实日志语料不存在：{REAL_LOG_ROOT}",
)

requires_golden = pytest.mark.skipif(
    not (GOLDEN_DIR / "fingerprints.json").is_file(),
    reason="黄金基线尚未生成，先跑 python tests/make_golden.py",
)
