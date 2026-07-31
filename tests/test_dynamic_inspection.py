"""动态巡检版本、VBS 与严格日志清单的关键回归。"""
from __future__ import annotations

import json
import sqlite3
from pathlib import Path

import pytest
from docx import Document
from fastapi import HTTPException

from core.inspection_docx import parse_template, validate_read_only_command, validate_snapshot
from core.vbs_generator import DIGEST_PLACEHOLDER, canonical_digest, generate_vbs


def _snapshot(name: str = 'R1"quoted') -> dict:
    return {
        "system_key": "IMS", "display_name": "IMS网络设备", "template": "template.docx",
        "devices": [{"order": 1, "name": name, "ip": "10.0.0.1", "driver": "huawei_vrp", "commands": [
            {"command": "display version", "timeout_seconds": 30, "result_cell": {"table": 0, "row": 4, "column": 2}},
            {"command": "display device", "timeout_seconds": 60, "result_cell": {"table": 0, "row": 5, "column": 2}},
        ]}], "non_command_rules": [],
    }


def test_template_parser_extracts_device_commands_and_mapping(tmp_path: Path):
    path = tmp_path / "inspection.docx"; doc = Document(); table = doc.add_table(rows=6, cols=3)
    table.cell(0, 0).text = "设备名称: IMS-R1"; table.cell(1, 0).text = "IP: 10.0.0.1"
    table.cell(4, 1).text = "display version"; table.cell(5, 1).text = "display device"; doc.save(path)
    result = parse_template(path, system_key="IMS", display_name="IMS网络设备")
    assert result["devices"][0]["name"] == "IMS-R1"
    assert result["devices"][0]["ip"] == "10.0.0.1"
    assert result["devices"][0]["commands"][1]["result_cell"] == {"table": 0, "row": 5, "column": 2}


def test_template_parser_uses_headers_and_skips_device_inventory(tmp_path: Path):
    path = tmp_path / "inspection.docx"; doc = Document()
    inventory = doc.add_table(rows=2, cols=4)
    for column, value in enumerate(("区域", "网元类型", "网元名称", "网元IP")):
        inventory.cell(0, column).text = value
    for column, value in enumerate(("拉萨", "CE8861", "JZ-5GToB-EOR-01", "10.237.1.98")):
        inventory.cell(1, column).text = value
    table = doc.add_table(rows=4, cols=6)
    for column, value in enumerate(("设备型号", "CE8861", "设备名称", "JZ-5GToB-EOR-01", "设备IP", "10.237.1.98")):
        table.cell(0, column).text = value
    for column, value in enumerate(("巡检类别", "巡检项", "巡检命令", "巡检结果", "巡检状态", "备注")):
        table.cell(1, column).text = value
    table.cell(2, 1).text = "配置备份"; table.cell(2, 2).text = "Display current-configuration"
    table.cell(3, 1).text = "版本核对"; table.cell(3, 2).text = "Display version"; table.cell(3, 3).text = "<LW-5GToB-EOR-01>Display version\nVersion 1"
    doc.save(path)

    result = parse_template(path, system_key="5GTOB", display_name="5G ToB")

    assert len(result["devices"]) == 1
    assert result["devices"][0]["name"] == "LW-5GToB-EOR-01"
    assert result["devices"][0]["ip"] == "10.237.1.98"
    assert [item["command"] for item in result["devices"][0]["commands"]] == ["Display current-configuration", "Display version"]
    assert result["devices"][0]["commands"][0]["result_cell"] == {"table": 1, "row": 2, "column": 3}


@pytest.mark.parametrize("command", ["system-view", "display version; reboot", "save", "undo interface x", "show run | delete"])
def test_dangerous_commands_are_blocked(command: str):
    with pytest.raises(HTTPException): validate_read_only_command(command)


def test_snapshot_rejects_duplicate_name_and_ip():
    snapshot = _snapshot("R1"); snapshot["devices"].append({**snapshot["devices"][0], "order": 2})
    with pytest.raises(HTTPException) as error: validate_snapshot(snapshot)
    assert error.value.status_code == 409


def test_vbs_escapes_dynamic_values_and_has_stable_canonical_digest():
    script, digest = generate_vbs("IMS", 3, validate_snapshot(_snapshot()))
    assert 'R1""quoted' in script
    assert "display version" in script and "Const SYSTEM_VERSION = 3" in script
    assert digest in script
    assert canonical_digest(script) == digest
    assert DIGEST_PLACEHOLDER not in script
    assert 'logsPath = fso.BuildPath(basePath, "logs")' in script
    assert 'datePath = fso.BuildPath(logsPath, DateFolder(Date))' in script
    assert 'BuildLogRoot = fso.BuildPath(datePath, SYSTEM_KEY)' in script
    assert 'SendCommandAndWait screenObj, prompt, "screen-length 0 temporary", 10, 1' in script
    assert "WaitForLoginPrompt(screenObj, 10)" in script
    assert "WaitForPromptOrHandleMore" in script
    assert "HasVisiblePasswordDecisionPrompt" in script
    assert "tabObj.Session.Log True" in script
    assert "tabObj.Session.Log True, True" not in script
    assert 'SYSTEM_KEY & "_summary_" & runStamp & ".log"' in script
    assert '"device_not_connected"' in script
    assert '"logging_failed"' in script
    assert "DeviceAlias(name)" in script


def _write_valid_batch(root: Path, *, digest: str, status: str = "success", include_second: bool = True):
    rows = [
        "system_key\tsystem_version\tscript_sha256\tdevice_name\tip\tlog_file\tcommand_index\tcommand\ttimeout_seconds\tstatus",
        f"IMS\t3\t{digest}\tR1\t10.0.0.1\tR1_10.0.0.1.log\t1\tdisplay version\t30\t{status}",
    ]
    if include_second: rows.append(f"IMS\t3\t{digest}\tR1\t10.0.0.1\tR1_10.0.0.1.log\t2\tdisplay device\t60\tsuccess")
    (root / "inspection-manifest.tsv").write_text("\n".join(rows), encoding="utf-8")
    (root / "R1_10.0.0.1.log").write_text("<R1>display version\nVersion 1\n<R1>display device\nDevice normal\n<R1>", encoding="utf-8")


def test_manifest_requires_every_command_and_digest(tmp_path: Path):
    from backend.log_batches import validate

    version = {"system_key": "IMS", "version": 3, "vbs_sha256": "a" * 64, "config_json": json.dumps(_snapshot("R1"), ensure_ascii=False)}
    _write_valid_batch(tmp_path, digest="a" * 64)
    report, _ = validate(tmp_path, version); assert report["valid"] is True
    _write_valid_batch(tmp_path, digest="b" * 64, include_second=False)
    report, _ = validate(tmp_path, version)
    assert report["valid"] is False
    assert any("摘要不符" in issue for issue in report["issues"])
    assert any("缺命令" in issue for issue in report["issues"])


def test_manifest_blocks_timeout_and_cli_error(tmp_path: Path):
    from backend.log_batches import validate

    version = {"system_key": "IMS", "version": 3, "vbs_sha256": "a" * 64, "config_json": json.dumps(_snapshot("R1"), ensure_ascii=False)}
    _write_valid_batch(tmp_path, digest="a" * 64, status="timeout")
    log = tmp_path / "R1_10.0.0.1.log"; log.write_text(log.read_text(encoding="utf-8") + "\nError: unknown command", encoding="utf-8")
    report, _ = validate(tmp_path, version)
    assert not report["valid"]
    assert any("timeout" in issue for issue in report["issues"])
    assert any("CLI 错误" in issue for issue in report["issues"])


def test_report_index_prefers_job_report_date(tmp_path: Path):
    from backend.reports import rebuild_report_file_index

    database = tmp_path / "app.db"; report = tmp_path / "report.docx"; Document().save(report)
    conn = sqlite3.connect(database); conn.row_factory = sqlite3.Row
    conn.executescript("""CREATE TABLE users(id INTEGER PRIMARY KEY, username TEXT); CREATE TABLE jobs(id TEXT PRIMARY KEY,user_id INTEGER,created_at TEXT,report_date TEXT,generated_files TEXT); CREATE TABLE report_files(id INTEGER PRIMARY KEY,job_id TEXT,user_id INTEGER,username TEXT,report_date TEXT,file_name TEXT,file_path TEXT UNIQUE,file_size INTEGER,modified_at TEXT,created_at TEXT);""")
    conn.execute("INSERT INTO users VALUES (1,'admin')"); conn.execute("INSERT INTO jobs VALUES ('j1',1,'2026-01-01T00:00:00+08:00','2026-07-24',?)", (json.dumps([str(report)]),)); conn.commit(); conn.close()
    def connect():
        value=sqlite3.connect(database); value.row_factory=sqlite3.Row; return value
    from datetime import timezone
    rebuild_report_file_index(db_connect=connect, local_tz=timezone.utc)
    conn=connect(); indexed=conn.execute("SELECT report_date FROM report_files").fetchone()[0]; conn.close()
    assert indexed == "2026-07-24"
